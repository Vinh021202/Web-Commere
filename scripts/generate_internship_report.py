from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path(r"H:\Web-Commere\20001375_TranQuocVinh_BCTTDN_Web-Commere_FINAL2.docx")


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_start = OxmlElement("w:fldChar")
    fld_start.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Mục lục sẽ được cập nhật khi mở file trong Word và chọn Update Field."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_start)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(txt)
    run._r.append(fld_end)


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(0.9)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.styles["Normal"].font.size = Pt(13)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(13)

    def center(text, size=13, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
        return p

    def para(text=""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        set_run_font(r)
        return p

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        set_run_font(r)
        return p

    def heading(text, level=1):
        p = doc.add_paragraph(style=f"Heading {level}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, size=15 if level == 1 else 14 if level == 2 else 13, bold=True)
        return p

    # Trang bìa
    center("TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP. HỒ CHÍ MINH", 14, True)
    center("KHOA CÔNG NGHỆ THÔNG TIN", 14, True)
    center("***********", 13, True)
    for _ in range(3):
        center("")
    center("BÁO CÁO THỰC TẬP DOANH NGHIỆP", 18, True)
    center("Ngành Công nghệ thông tin / Kỹ thuật phần mềm", 14)
    for _ in range(2):
        center("")
    center("Đề tài:", 14, True)
    center("XÂY DỰNG VÀ PHÁT TRIỂN HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ WEB-COMMERE", 15, True)
    for _ in range(3):
        center("")
    center("Giảng viên hướng dẫn: ........................................................", 13)
    center("Người hướng dẫn tại doanh nghiệp: ...................................", 13)
    center("Sinh viên thực hiện: Trần Quốc Vịnh", 13)
    center("Mã sinh viên: 20001375", 13)
    center("Lớp: ................................................", 13)
    center("Đơn vị thực tập: ...............................................................", 13)
    for _ in range(5):
        center("")
    center("TP. Hồ Chí Minh, năm 2026", 13)

    doc.add_page_break()

    heading("LỜI CẢM ƠN", 1)
    para(
        "Trước tiên, em xin gửi lời cảm ơn chân thành đến Quý Thầy Cô Khoa Công nghệ Thông tin – "
        "Trường Đại học Công nghiệp TP. Hồ Chí Minh đã tận tình giảng dạy, truyền đạt cho em những "
        "kiến thức chuyên môn và kỹ năng cần thiết trong suốt quá trình học tập tại trường. Những kiến "
        "thức đó chính là nền tảng quan trọng giúp em có thể tiếp cận và hoàn thành kỳ thực tập doanh "
        "nghiệp một cách nghiêm túc, hiệu quả."
    )
    para(
        "Em xin chân thành cảm ơn giảng viên hướng dẫn đã luôn theo sát, góp ý và định hướng để em hoàn "
        "thiện báo cáo này. Đồng thời, em cũng xin cảm ơn đơn vị thực tập và người hướng dẫn tại doanh "
        "nghiệp đã tạo điều kiện cho em được tiếp cận môi trường làm việc thực tế, được tham gia tìm hiểu, "
        "phân tích và phát triển dự án Web-Commere – một hệ thống thương mại điện tử full-stack với nhiều "
        "nghiệp vụ thực tiễn."
    )
    para(
        "Trong suốt quá trình thực tập, em đã có cơ hội rèn luyện thêm về tư duy lập trình, quy trình làm "
        "việc nhóm, kỹ năng sử dụng Git, cách đọc hiểu hệ thống lớn và nhận thức rõ hơn về vai trò của bảo "
        "mật, tài liệu hóa và tính kỷ luật trong phát triển phần mềm. Đây là những kinh nghiệm quý báu giúp "
        "em trưởng thành hơn trong học tập cũng như định hướng nghề nghiệp sau này."
    )
    para("Em xin chân thành cảm ơn!")

    doc.add_page_break()

    center("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 13, True)
    center("Độc lập – Tự do – Hạnh phúc", 13, True)
    center("")
    center("NHẬN XÉT CỦA ĐƠN VỊ THỰC TẬP", 14, True)
    para("Họ và tên sinh viên: Trần Quốc Vịnh")
    para("Mã sinh viên: 20001375")
    para("Nhận xét chung:")
    for _ in range(12):
        para("............................................................................................................................")
    para("")
    para("Xác nhận của đơn vị thực tập")

    doc.add_page_break()

    center("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 14, True)
    for _ in range(15):
        para("............................................................................................................................")

    doc.add_page_break()

    heading("MỤC LỤC", 1)
    add_toc(doc)

    doc.add_page_break()

    heading("CHƯƠNG 1. GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP", 1)
    heading("1.1. Thông tin về đơn vị thực tập", 2)
    para(
        "Trong thời gian thực tập doanh nghiệp, em có cơ hội tiếp cận môi trường làm việc thực tế tại đơn vị "
        "thực tập trong lĩnh vực công nghệ thông tin và phát triển phần mềm. Tại đây, em được tham gia tìm "
        "hiểu quy trình xây dựng sản phẩm số, đặc biệt là các dự án web ứng dụng phục vụ thương mại điện tử "
        "và quản lý dữ liệu doanh nghiệp."
    )
    para(
        "Đơn vị thực tập hướng đến mô hình làm việc hiện đại, chú trọng vào việc tổ chức source code rõ ràng, "
        "quản lý công việc theo từng module và đề cao khả năng tự học, tự nghiên cứu của thực tập sinh. Đây "
        "là môi trường phù hợp để sinh viên có thể vận dụng kiến thức đã học ở trường vào thực tế và hiểu rõ "
        "hơn cách một dự án phần mềm được vận hành trong doanh nghiệp."
    )
    heading("1.2. Thông tin về vị trí sinh viên tham gia thực tập", 2)
    para(
        "Trong kỳ thực tập này, em tham gia với vai trò thực tập sinh phát triển phần mềm, tập trung vào việc "
        "đọc hiểu hệ thống, hỗ trợ cập nhật tài liệu dự án, phân tích chức năng, chỉnh sửa giao diện và chuẩn "
        "hóa cấu hình môi trường cho dự án Web-Commere. Vị trí này yêu cầu sinh viên có kiến thức nền tảng "
        "về lập trình web, hiểu cách frontend và backend phối hợp với nhau, đồng thời có tinh thần cẩn thận "
        "khi thao tác với Git và dữ liệu cấu hình."
    )
    bullet("Nắm được cấu trúc dự án web full-stack gồm client, admin và server.")
    bullet("Có khả năng đọc hiểu React, Express, MongoDB và các API RESTful.")
    bullet("Biết sử dụng Git để kiểm soát thay đổi mã nguồn.")
    bullet("Rèn luyện kỹ năng viết tài liệu kỹ thuật và mô tả chức năng hệ thống.")

    heading(
        "CHƯƠNG 2. VẤN ĐỀ MÀ SINH VIÊN THAM GIA GIẢI QUYẾT/THỰC HIỆN TẠI ĐƠN VỊ/DOANH NGHIỆP THỰC TẬP",
        1,
    )
    heading("2.1. Tóm tắt vấn đề mà sinh viên tham gia giải quyết/thực hiện tại đơn vị/doanh nghiệp thực tập", 2)
    para(
        "Trong thời gian thực tập, em tham gia tìm hiểu và hỗ trợ phát triển dự án Web-Commere – một hệ thống "
        "thương mại điện tử full-stack gồm giao diện khách hàng, trang quản trị và backend API. Dự án hướng đến "
        "việc xây dựng một nền tảng bán hàng trực tuyến có thể đáp ứng các nhu cầu cơ bản như quản lý sản phẩm, "
        "danh mục, giỏ hàng, wishlist, địa chỉ giao hàng, đơn hàng, thanh toán, banner, blog và quản trị người dùng."
    )
    para(
        "Các công việc em tham gia chủ yếu xoay quanh việc đọc hiểu toàn bộ source code, phân tích kiến trúc hệ "
        "thống, tìm hiểu luồng dữ liệu giữa các thành phần client/admin/server, chỉnh sửa một số giao diện khu vực "
        "tài khoản người dùng, viết lại README mô tả dự án và đặc biệt là chuẩn hóa cấu hình môi trường bằng các "
        "file .env.example nhằm tránh rò rỉ thông tin nhạy cảm khi làm việc với Git và GitHub."
    )

    heading("2.2. Tiến độ thực hiện công việc (các mốc thời gian thực hiện)", 2)
    para(
        "Trong suốt thời gian thực tập, em triển khai công việc theo từng giai đoạn cụ thể, từ làm quen với môi "
        "trường phát triển, phân tích dự án, tìm hiểu nghiệp vụ cho đến chỉnh sửa giao diện, tài liệu hóa hệ thống "
        "và chuẩn hóa cấu hình. Mỗi giai đoạn đều có mục tiêu riêng và giúp em hiểu sâu hơn về dự án cũng như quy "
        "trình làm việc thực tế trong doanh nghiệp."
    )

    doc.add_page_break()
    center("TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP HCM", 12, True)
    center("KHOA CÔNG NGHỆ THÔNG TIN", 12, True)
    center("")
    center("NHẬT KÝ THỰC TẬP", 14, True)
    center("")
    para("Tên công ty: ...............................................................")
    para("Địa chỉ: .....................................................................")
    para("Người hướng dẫn: ........................................................")
    para("Họ và tên: Trần Quốc Vịnh                MSSV: 20001375")
    para("Điện thoại: ....................................     Email: ....................................")

    table = doc.add_table(rows=13, cols=4)
    table.style = "Table Grid"
    headers = ["Tuần", "Thời gian", "Nội dung công việc", "Kết quả/Ghi chú"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, bold=True)

    log_rows = [
        ("1", "Tuần 1", "Tìm hiểu môi trường, cấu trúc client/admin/server, cài đặt công cụ phát triển.", "Đã nắm cấu trúc tổng quan dự án."),
        ("2", "Tuần 2", "Phân tích các chức năng chính của hệ thống và luồng nghiệp vụ.", "Hiểu các module sản phẩm, người dùng, đơn hàng."),
        ("3", "Tuần 3", "Tìm hiểu luồng dữ liệu giữa frontend và backend thông qua API.", "Nắm được cơ chế gọi API và xác thực."),
        ("4", "Tuần 4", "Đọc module tài khoản người dùng và địa chỉ giao hàng.", "Chuẩn bị cho chỉnh sửa giao diện tài khoản."),
        ("5", "Tuần 5", "Điều chỉnh giao diện My Account và Address.", "Cải thiện hiển thị và nội dung tiếng Việt."),
        ("6", "Tuần 6", "Viết lại README mô tả dự án.", "Hoàn thiện README rõ ràng hơn."),
        ("7", "Tuần 7", "Chuẩn hóa file .env.example.", "Tách cấu hình mẫu khỏi secret thật."),
        ("8", "Tuần 8", "Rà soát Git và xử lý rủi ro lộ thông tin nhạy cảm.", "Đảm bảo chỉ push thay đổi an toàn."),
        ("9", "Tuần 9", "Nghiên cứu sâu backend và mô hình dữ liệu.", "Hiểu thêm về controller, route, model."),
        ("10", "Tuần 10", "Tổng hợp kiến trúc hệ thống và nội dung báo cáo.", "Hoàn thành khung nội dung báo cáo."),
        ("11", "Tuần 11", "Đối chiếu với đề cương mẫu và chỉnh sửa trình bày.", "Bố cục bám sát mẫu thực tập."),
        ("12", "Tuần 12", "Hoàn thiện báo cáo thực tập doanh nghiệp.", "Xuất file Word hoàn chỉnh."),
    ]

    for row_index, row_data in enumerate(log_rows, start=1):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index >= 2 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run)

    heading("2.3. Tóm tắt công việc thực hiện theo tiến độ (báo cáo hàng tuần).", 2)

    weekly_sections = [
        (
            "2.3.1 Cài đặt môi trường phát triển",
            "Trong tuần đầu tiên của kỳ thực tập, em tập trung tìm hiểu môi trường làm việc thực tế và chuẩn bị đầy "
            "đủ công cụ phục vụ cho quá trình tham gia dự án. Em tiến hành rà soát cấu trúc repository, xác định ba "
            "thành phần chính của hệ thống là client, admin và server, đồng thời cài đặt và kiểm tra các công cụ cần "
            "thiết như Node.js, npm, Visual Studio Code, Git và các extension hỗ trợ lập trình React/JavaScript. Bên "
            "cạnh đó, em cũng tìm hiểu cách dự án kết nối MongoDB, sử dụng Express cho backend và cách tổ chức source "
            "code theo từng module nghiệp vụ. Kết quả của giai đoạn này là em đã có môi trường phát triển ổn định và "
            "nắm được cấu trúc tổng quát của dự án Web-Commere để sẵn sàng cho các bước tiếp theo."
        ),
        (
            "2.3.2 Tham khảo giao diện của các website tương tự và tìm hiểu cấu trúc hiển thị của dự án",
            "Trong giai đoạn tiếp theo, em dành thời gian nghiên cứu giao diện của các website thương mại điện tử có "
            "chức năng tương đồng nhằm hiểu rõ hơn cách tổ chức bố cục, hiển thị sản phẩm và xây dựng trải nghiệm "
            "người dùng. Song song với việc tham khảo, em đọc các trang chính trong dự án hiện tại như trang chủ, "
            "trang danh sách sản phẩm, trang chi tiết sản phẩm, khu vực tài khoản cá nhân và giao diện quản trị để "
            "nắm được định hướng thiết kế mà hệ thống đang sử dụng. Quá trình này giúp em hiểu cách frontend được "
            "xây dựng bằng React, cách các component được tách nhỏ và cách giao diện được liên kết với API backend."
        ),
        (
            "2.3.3 Phân tích trang sản phẩm và luồng dữ liệu API RESTful",
            "Ở tuần này, em tập trung phân tích sâu hơn module sản phẩm – đây là phần trung tâm của hệ thống thương "
            "mại điện tử. Em tìm hiểu cách frontend lấy dữ liệu sản phẩm, hiển thị danh sách, phân trang và liên kết "
            "đến trang chi tiết. Đồng thời, em đọc các route và controller phía backend để hiểu cách API RESTful được "
            "tổ chức cho các thao tác như lấy tất cả sản phẩm, lấy sản phẩm theo danh mục, lọc theo giá, sắp xếp, tìm "
            "kiếm và lấy chi tiết theo ID. Qua đó, em hiểu rõ hơn mối liên hệ giữa dữ liệu lưu trong MongoDB và cách "
            "hiển thị ra giao diện người dùng."
        ),
        (
            "2.3.4 Tìm hiểu giỏ hàng, địa chỉ giao hàng và khu vực tài khoản người dùng",
            "Tiếp theo, em chuyển sang nghiên cứu nhóm chức năng liên quan trực tiếp đến quá trình mua hàng của khách "
            "hàng. Em tìm hiểu cách hệ thống quản lý giỏ hàng, cách thêm/xóa/cập nhật số lượng sản phẩm và cách dữ "
            "liệu giỏ hàng được đồng bộ với tài khoản người dùng. Ngoài ra, em cũng phân tích chức năng quản lý địa "
            "chỉ giao hàng, cập nhật hồ sơ cá nhân và đổi mật khẩu trong khu vực My Account. Việc nắm được những phần "
            "này giúp em chuẩn bị tốt hơn cho công việc chỉnh sửa giao diện và nội dung hiển thị ở các tuần sau."
        ),
        (
            "2.3.5 Điều chỉnh giao diện trang My Account và Address",
            "Sau khi đã nắm được chức năng của khu vực tài khoản người dùng, em tham gia chỉnh sửa một số chi tiết giao "
            "diện tại các trang My Account, Address và form thêm địa chỉ. Công việc bao gồm rà soát khoảng cách hiển "
            "thị, cải thiện phần tiêu đề và nội dung tiếng Việt, đồng thời làm cho bố cục trực quan và thống nhất hơn "
            "với giao diện chung của website. Việc chỉnh sửa này tuy không quá lớn về mặt logic nhưng có ý nghĩa trong "
            "việc nâng cao trải nghiệm người dùng và làm cho phần tài khoản cá nhân trở nên hoàn chỉnh hơn."
        ),
        (
            "2.3.6 Viết lại README mô tả dự án",
            "Trong tuần này, em tham gia viết lại README của dự án theo hướng đầy đủ, rõ ràng và phù hợp hơn với một "
            "repository phát triển phần mềm thực tế. Nội dung tài liệu được cập nhật bao gồm phần giới thiệu dự án, "
            "mục tiêu hệ thống, các tính năng nổi bật, kiến trúc tổng thể, hướng dẫn cài đặt, cách chạy, cấu trúc thư "
            "mục, công nghệ sử dụng, hướng đóng góp và license. Đây là một công việc quan trọng vì README đóng vai trò "
            "là tài liệu đầu vào giúp người khác nhanh chóng hiểu dự án và biết cách sử dụng source code."
        ),
        (
            "2.3.7 Chuẩn hóa cấu hình môi trường bằng file .env.example",
            "Một nội dung rất quan trọng mà em được tiếp cận là chuẩn hóa cấu hình môi trường của hệ thống. Em tiến hành "
            "tạo và cập nhật các file .env.example cho client, admin và server dựa trên các biến môi trường mà dự án "
            "đang sử dụng. Quá trình này giúp tách biệt phần cấu hình thật với phần cấu hình mẫu, vừa hỗ trợ người khác "
            "dễ triển khai dự án, vừa tránh đưa những thông tin nhạy cảm như API key, secret token hay thông tin dịch "
            "vụ bên thứ ba lên remote repository."
        ),
        (
            "2.3.8 Làm việc với Git và xử lý rủi ro lộ thông tin nhạy cảm",
            "Trong quá trình thao tác với Git, em gặp một tình huống thực tế liên quan đến việc các file .env có chứa "
            "secret thật có nguy cơ bị đẩy lên GitHub. Từ đó, em tham gia rà soát các commit local, phân biệt thay đổi "
            "an toàn và thay đổi nhạy cảm, sau đó xây dựng một branch sạch để chỉ giữ lại các phần được phép công khai "
            "như README, giao diện và file .env.example. Qua tình huống này, em rút ra được bài học quan trọng về kiểm "
            "tra lịch sử commit, review file trước khi push và tư duy bảo mật trong quản lý mã nguồn."
        ),
        (
            "2.3.9 Rà soát các module backend và mô hình dữ liệu",
            "Tuần này em tiếp tục nghiên cứu sâu hơn ở backend để hiểu cách hệ thống được tổ chức theo từng domain "
            "nghiệp vụ. Em đọc các controller, route và model thuộc những module như user, product, category, cart, "
            "myList, address, order, bannerV1, blog, homeSlider và chat. Bên cạnh đó, em cũng tìm hiểu cách Mongoose "
            "schema được sử dụng để mô tả dữ liệu và mối liên hệ giữa các thực thể như người dùng, địa chỉ, sản phẩm "
            "và đơn hàng. Việc này giúp em có cái nhìn toàn diện hơn về kiến trúc backend của dự án."
        ),
        (
            "2.3.10 Tổng hợp kiến trúc hệ thống và biên soạn nội dung báo cáo",
            "Sau khi đã có đủ bức tranh tổng thể về hệ thống, em bắt đầu tổng hợp các nội dung phục vụ báo cáo thực tập "
            "doanh nghiệp. Em hệ thống lại kiến trúc dự án, các công nghệ đang sử dụng, những chức năng chính của hệ "
            "thống, các công việc thực hiện theo tiến độ và những kết quả đạt được trong quá trình tham gia dự án. Đây "
            "là giai đoạn chuyển từ quá trình quan sát và chỉnh sửa kỹ thuật sang quá trình tài liệu hóa và trình bày "
            "một cách học thuật, rõ ràng hơn."
        ),
        (
            "2.3.11 Hoàn thiện nội dung trình bày và đối chiếu với đề cương mẫu",
            "Ở giai đoạn gần cuối, em tiến hành rà soát lại toàn bộ cấu trúc báo cáo để bảo đảm cách trình bày bám sát "
            "đề cương thực tập doanh nghiệp. Em chỉnh sửa lại tên các mục, thêm những phần còn thiếu như tiến độ thực "
            "hiện công việc, nhật ký thực tập và phần tóm tắt theo tuần, đồng thời viết lại các nội dung theo đúng văn "
            "phong của mẫu báo cáo. Công việc này giúp báo cáo vừa phản ánh đúng nội dung em đã thực hiện, vừa đáp ứng "
            "được yêu cầu hình thức của nhà trường."
        ),
        (
            "2.3.12 Hoàn tất báo cáo thực tập doanh nghiệp",
            "Trong tuần cuối cùng, em hoàn thiện bản báo cáo thực tập doanh nghiệp và kiểm tra lại toàn bộ nội dung, "
            "định dạng, chương mục và phần thông tin cá nhân. Đồng thời, em rà soát kỹ lại file Word để bảo đảm không "
            "còn lỗi font chữ, lỗi tiếng Việt hay thiếu sót ở các phần nhận xét, mục lục và các chương nội dung. Kết "
            "quả đạt được là một bản báo cáo hoàn chỉnh, có cấu trúc rõ ràng, phản ánh trung thực quá trình thực tập "
            "và bám sát đề cương mẫu mà em đã tham khảo."
        ),
    ]

    for title, content in weekly_sections:
        heading(title, 3)
        para(content)

    heading("CHƯƠNG 3. NHẬN XÉT, ĐÁNH GIÁ CHƯƠNG TRÌNH THỰC TẬP", 1)
    para(
        "Trong suốt quá trình thực tập doanh nghiệp, em có cơ hội tiếp cận một dự án phần mềm thực tế với cấu trúc "
        "tương đối đầy đủ và rõ ràng. Việc tham gia phân tích, đọc hiểu và chỉnh sửa dự án Web-Commere giúp em nhận "
        "ra sự khác biệt giữa bài tập học thuật và một hệ thống ứng dụng thật, nơi mọi thay đổi đều cần cân nhắc đến "
        "kiến trúc, bảo trì, khả năng mở rộng và tính an toàn của dữ liệu."
    )
    para(
        "Kỳ thực tập giúp em rèn luyện nhiều kỹ năng quan trọng như đọc hiểu source code lớn, xác định luồng dữ liệu "
        "giữa frontend và backend, thao tác Git một cách cẩn trọng, viết tài liệu mô tả dự án và nhận thức rõ vai trò "
        "của file cấu hình môi trường trong bảo mật hệ thống. Không chỉ dừng ở khía cạnh kỹ thuật, em còn học được cách "
        "làm việc có trình tự, biết ưu tiên vấn đề cần xử lý và cẩn thận hơn trong mọi thao tác với mã nguồn."
    )
    heading("3.1. Những kiến thức và kỹ năng đạt được", 2)
    bullet("Hiểu rõ hơn về kiến trúc một dự án web full-stack gồm client, admin và server.")
    bullet("Củng cố kiến thức về React, Express, MongoDB và API RESTful.")
    bullet("Biết cách tài liệu hóa dự án thông qua README và báo cáo kỹ thuật.")
    bullet("Nâng cao nhận thức về an toàn thông tin khi làm việc với .env, token và secret key.")
    bullet("Rèn luyện kỹ năng sử dụng Git, đặc biệt là kiểm tra commit và push an toàn.")
    heading("3.2. Những khó khăn gặp phải", 2)
    bullet("Dự án có nhiều module nên việc nắm toàn bộ cấu trúc ban đầu tương đối mất thời gian.")
    bullet("Một số file cũ có nội dung hiển thị chưa đồng nhất hoặc tồn tại lỗi ký tự, cần rà soát kỹ khi chỉnh sửa.")
    bullet("Việc xử lý commit có liên quan đến file cấu hình nhạy cảm đòi hỏi sự cẩn trọng cao.")
    heading("3.3. Đánh giá chung", 2)
    para(
        "Nhìn chung, chương trình thực tập mang lại giá trị thực tiễn cao. Em không chỉ học được thêm về công nghệ "
        "mà còn hiểu hơn về quy trình, kỷ luật và trách nhiệm trong phát triển phần mềm. Đây là nền tảng tốt để em "
        "tiếp tục nâng cao năng lực chuyên môn trong các học phần tiếp theo và trong công việc sau khi ra trường."
    )

    heading("CHƯƠNG 4. KẾT LUẬN", 1)
    heading("4.1. Kết luận", 2)
    para(
        "Kỳ thực tập doanh nghiệp là cơ hội quan trọng giúp em áp dụng kiến thức đã học vào môi trường thực tế, đồng "
        "thời hiểu rõ hơn cách một dự án công nghệ được tổ chức, phát triển và bảo trì. Thông qua dự án Web-Commere, "
        "em đã tiếp cận được nhiều khía cạnh của phát triển phần mềm như frontend, backend, tài liệu hóa, quản lý cấu "
        "hình và thao tác Git an toàn."
    )
    para(
        "Những trải nghiệm trong quá trình thực tập giúp em tự tin hơn về mặt chuyên môn, đặc biệt là khả năng tự học, "
        "khả năng đọc hiểu hệ thống và tư duy giải quyết vấn đề. Đây là hành trang quan trọng để em tiếp tục hoàn thiện "
        "bản thân và chuẩn bị tốt hơn cho công việc trong lĩnh vực công nghệ thông tin."
    )
    heading("4.2. Kiến nghị và hướng phát triển", 2)
    bullet("Tiếp tục mở rộng các tính năng của hệ thống như tối ưu tìm kiếm, lọc sản phẩm và quản lý đơn hàng.")
    bullet("Chuẩn hóa sâu hơn tài liệu kỹ thuật, quy trình deploy và quy trình quản lý biến môi trường.")
    bullet("Tăng cường các biện pháp bảo mật và kiểm thử để hệ thống ổn định hơn khi triển khai thực tế.")
    bullet("Hoàn thiện thêm phần giao diện người dùng để nâng cao trải nghiệm mua sắm trên nhiều thiết bị.")

    heading("TÀI LIỆU THAM KHẢO", 1)
    bullet("Tài liệu chính thức của React: https://react.dev/")
    bullet("Tài liệu chính thức của Express: https://expressjs.com/")
    bullet("Tài liệu chính thức của MongoDB: https://www.mongodb.com/")
    bullet("Tài liệu chính thức của Vite: https://vitejs.dev/")
    bullet("Source code dự án Web-Commere được sử dụng trong quá trình thực tập.")

    for sec in doc.sections:
        add_page_number(sec)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
